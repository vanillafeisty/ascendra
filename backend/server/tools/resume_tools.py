"""
CareerPilot AI - Resume Builder
Builds waterproof, ATS-optimized resumes from user profile.
Tailors resume to specific job descriptions.
Outputs DOCX and PDF.
"""

import re
import json
from pathlib import Path
from datetime import datetime
from typing import Optional

from config import RESUMES_DIR, load_user_profile, GROQ_API_KEY


# ─────────────────────────────────────────────────────────────────────────────
# ATS KEYWORD EXTRACTOR
# ─────────────────────────────────────────────────────────────────────────────

def extract_jd_keywords(job_description: str) -> list[str]:
    """
    Extract important keywords from a job description for ATS optimization.
    """
    # Tech keywords pattern
    tech_pattern = r'\b(?:Python|Java|JavaScript|TypeScript|React|Angular|Vue|Node\.?js|'
    tech_pattern += r'SQL|NoSQL|MongoDB|PostgreSQL|MySQL|Redis|Elasticsearch|'
    tech_pattern += r'AWS|Azure|GCP|Docker|Kubernetes|CI/CD|Git|REST|GraphQL|'
    tech_pattern += r'Machine Learning|Deep Learning|NLP|TensorFlow|PyTorch|'
    tech_pattern += r'Spring Boot|Django|FastAPI|Flask|Express|'
    tech_pattern += r'Agile|Scrum|DevOps|Microservices|API|Cloud|Linux|'
    tech_pattern += r'HTML|CSS|SASS|Tailwind|Bootstrap|Figma|'
    tech_pattern += r'C\+\+|C#|Go|Rust|Kotlin|Swift|PHP|Ruby)\b'

    tech_keywords = re.findall(tech_pattern, job_description, re.IGNORECASE)

    # Soft skill keywords
    soft_skills = ["leadership", "communication", "teamwork", "problem-solving",
                   "analytical", "collaboration", "initiative", "adaptable",
                   "detail-oriented", "time management"]
    soft_found = [s for s in soft_skills if s.lower() in job_description.lower()]

    # Combine and deduplicate
    all_keywords = list(dict.fromkeys([k.lower() for k in tech_keywords] + soft_found))
    return all_keywords


def score_resume_against_jd(resume_text: str, job_description: str) -> dict:
    """
    Score a resume against a job description (0-100).
    Returns score, matched keywords, and missing keywords.
    """
    jd_keywords = extract_jd_keywords(job_description)
    resume_lower = resume_text.lower()

    matched = [kw for kw in jd_keywords if kw in resume_lower]
    missing = [kw for kw in jd_keywords if kw not in resume_lower]

    score = int((len(matched) / len(jd_keywords)) * 100) if jd_keywords else 70

    return {
        "score": score,
        "grade": "A" if score >= 85 else "B" if score >= 70 else "C" if score >= 55 else "D",
        "matched_keywords": matched,
        "missing_keywords": missing,
        "total_jd_keywords": len(jd_keywords),
        "recommendation": (
            "Strong match! Submit with confidence." if score >= 85 else
            f"Add these missing keywords naturally: {', '.join(missing[:5])}" if missing else
            "Good match. Review for any other gaps."
        )
    }


# ─────────────────────────────────────────────────────────────────────────────
# RESUME BUILDER (DOCX)
# ─────────────────────────────────────────────────────────────────────────────

def build_resume_docx(profile: dict, output_filename: str = None) -> str:
    """
    Build an ATS-optimized DOCX resume from user profile dict.
    Returns path to the generated DOCX file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ── Styles ────────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def add_heading(text: str, level: int = 1, color: tuple = (31, 73, 125)):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11 if level == 1 else 10)
        run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        # Bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_bullet(text: str, indent: float = 0.25):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(text)
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)
        return p

    # ── HEADER ────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(profile.get("name", "Your Name").upper())
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(31, 73, 125)

    # Contact info line
    contact_parts = []
    if profile.get("phone"):      contact_parts.append(profile["phone"])
    if profile.get("email"):      contact_parts.append(profile["email"])
    if profile.get("location"):   contact_parts.append(profile["location"])
    if profile.get("linkedin_url"): contact_parts.append(profile["linkedin_url"])
    if profile.get("github_url"): contact_parts.append(profile["github_url"])
    if profile.get("portfolio_url"): contact_parts.append(profile["portfolio_url"])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_para.paragraph_format.space_after = Pt(4)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    if profile.get("summary"):
        add_heading("Professional Summary")
        p = doc.add_paragraph()
        p.add_run(profile["summary"])
        p.paragraph_format.space_after = Pt(2)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    if profile.get("skills"):
        add_heading("Technical Skills")
        # Group skills in rows of 6
        skills = profile["skills"]
        for i in range(0, len(skills), 6):
            chunk = skills[i:i+6]
            p = doc.add_paragraph()
            p.add_run("• " + "  •  ".join(chunk))
            p.paragraph_format.space_after = Pt(1)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    if profile.get("experience"):
        add_heading("Professional Experience")
        for exp in profile["experience"]:
            p = doc.add_paragraph()
            role_run = p.add_run(exp.get("role", ""))
            role_run.bold = True
            role_run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(0)

            p2 = doc.add_paragraph()
            company_run = p2.add_run(exp.get("company", ""))
            company_run.bold = True
            company_run.font.color.rgb = RGBColor(89, 89, 89)
            if exp.get("duration"):
                dur_run = p2.add_run(f"  |  {exp['duration']}")
                dur_run.font.color.rgb = RGBColor(89, 89, 89)
            p2.paragraph_format.space_after = Pt(2)

            for bullet in exp.get("bullets", []):
                add_bullet(bullet)

    # ── PROJECTS ──────────────────────────────────────────────────────────────
    if profile.get("projects"):
        add_heading("Projects")
        for proj in profile["projects"]:
            p = doc.add_paragraph()
            name_run = p.add_run(proj.get("name", ""))
            name_run.bold = True
            if proj.get("link"):
                link_run = p.add_run(f"  |  {proj['link']}")
                link_run.font.color.rgb = RGBColor(0, 112, 192)
            p.paragraph_format.space_after = Pt(1)

            if proj.get("tech_stack"):
                tech_p = doc.add_paragraph()
                tech_run = tech_p.add_run("Stack: " + ", ".join(proj["tech_stack"]))
                tech_run.italic = True
                tech_run.font.size = Pt(9)
                tech_p.paragraph_format.space_after = Pt(1)

            if proj.get("description"):
                add_bullet(proj["description"])

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    if profile.get("certifications"):
        add_heading("Certifications")
        for cert in profile["certifications"]:
            add_bullet(cert)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if profile.get("education"):
        add_heading("Education")
        for edu in profile["education"]:
            p = doc.add_paragraph()
            deg_run = p.add_run(edu.get("degree", ""))
            deg_run.bold = True
            p.paragraph_format.space_after = Pt(0)

            p2 = doc.add_paragraph()
            p2.add_run(edu.get("institution", ""))
            if edu.get("year"):   p2.add_run(f"  |  {edu['year']}")
            if edu.get("cgpa"):   p2.add_run(f"  |  CGPA: {edu['cgpa']}")
            p2.paragraph_format.space_after = Pt(4)

    # ── Save ──────────────────────────────────────────────────────────────────
    if not output_filename:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"resume_{ts}.docx"

    output_path = RESUMES_DIR / output_filename
    doc.save(str(output_path))
    return str(output_path)


def build_resume(profile: Optional[dict] = None) -> dict:
    """
    Build a complete ATS-optimized resume. Returns file paths.
    """
    if profile is None:
        profile = load_user_profile()

    if not profile.get("name"):
        return {
            "success": False,
            "reason": "User profile not set up. Please provide your profile details first."
        }

    try:
        docx_path = build_resume_docx(profile)
        return {
            "success": True,
            "docx_path": docx_path,
            "message": f"Resume built successfully! Saved to: {docx_path}",
            "tip": "Review the resume and run resume_tailor with a specific JD to optimize it further."
        }
    except Exception as e:
        return {"success": False, "reason": str(e)}


def tailor_resume_to_jd(resume_path: str, job_description: str) -> dict:
    """
    Tailor an existing resume to a specific job description.
    Injects missing keywords and optimizes bullet points.
    Returns path to the tailored resume.
    """
    try:
        from docx import Document

        resume_path = Path(resume_path)
        if not resume_path.exists():
            return {"success": False, "reason": f"Resume not found: {resume_path}"}

        doc = Document(str(resume_path))
        full_text = "\n".join([p.text for p in doc.paragraphs])

        # Score before
        score_before = score_resume_against_jd(full_text, job_description)
        missing_kw = score_before["missing_keywords"]

        # Add skills section if keywords are missing
        if missing_kw:
            skills_addition = f"\n\nADDITIONAL RELEVANT SKILLS: {', '.join(missing_kw[:10])}"
            # Find and update skills paragraph
            for para in doc.paragraphs:
                if "Technical Skills" in para.text or "Skills" in para.text:
                    next_para = doc.paragraphs[doc.paragraphs.index(para) + 1]
                    next_para.runs[0].text += f"  •  {', '.join(missing_kw[:8])}"
                    break

        # Save tailored version
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        tailored_path = RESUMES_DIR / f"resume_tailored_{ts}.docx"
        doc.save(str(tailored_path))

        # Score after
        full_text_after = "\n".join([p.text for p in doc.paragraphs])
        score_after = score_resume_against_jd(full_text_after, job_description)

        return {
            "success": True,
            "tailored_path": str(tailored_path),
            "score_before": score_before["score"],
            "score_after": score_after["score"],
            "keywords_added": missing_kw[:8],
            "message": f"Resume tailored! ATS score improved from {score_before['score']} → {score_after['score']}"
        }
    except Exception as e:
        return {"success": False, "reason": str(e)}


def analyze_resume(resume_path: str) -> dict:
    """
    Analyze an uploaded resume and return detailed feedback.
    """
    try:
        from docx import Document

        rp = Path(resume_path)
        if not rp.exists():
            return {"success": False, "reason": f"File not found: {resume_path}"}

        if rp.suffix == ".docx":
            doc = Document(str(rp))
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            return {"success": False, "reason": "Only DOCX format supported for analysis. Convert PDF to DOCX first."}

        word_count = len(text.split())
        has_quantified = bool(re.search(r'\d+%|\d+ (years?|months?|people|users|projects?)', text, re.I))
        has_action_verbs = bool(re.search(r'\b(Led|Built|Developed|Designed|Improved|Increased|Reduced|Managed|Created|Delivered)\b', text))

        feedback = []
        if word_count < 300:
            feedback.append("⚠️ Resume is too short. Add more details to projects and experience.")
        if word_count > 900:
            feedback.append("⚠️ Resume may be too long for 0–3 years experience. Keep to 1 page.")
        if not has_quantified:
            feedback.append("⚠️ No quantified achievements found. Add numbers: '↑ performance by 40%', 'served 1000+ users'")
        if not has_action_verbs:
            feedback.append("⚠️ Start bullet points with strong action verbs: Built, Led, Designed, Optimized, etc.")
        if "summary" not in text.lower() and "objective" not in text.lower():
            feedback.append("⚠️ No professional summary found. Add a 2–3 sentence summary at the top.")

        return {
            "success": True,
            "word_count": word_count,
            "has_quantified_achievements": has_quantified,
            "has_action_verbs": has_action_verbs,
            "issues_found": len(feedback),
            "feedback": feedback if feedback else ["✅ Resume looks solid! Tailor it to a specific JD for best results."],
            "text_preview": text[:500] + "..."
        }
    except Exception as e:
        return {"success": False, "reason": str(e)}
